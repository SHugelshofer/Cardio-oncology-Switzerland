#!/usr/bin/env python3
"""
Convert Markdown files to Word documents for grant application submission.

This script converts markdown files from the grant_application directory 
into professional Word documents (.docx) suitable for submission to 
Fondation Emma Muschamp.

Usage:
    python3 convert_to_word.py [--all | file1.md file2.md ...]
    
Examples:
    # Convert all markdown files
    python3 convert_to_word.py --all
    
    # Convert specific files
    python3 convert_to_word.py grant_application/resume_projet.md grant_application/budget_detaille.md
"""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx not installed. Please run: pip install python-docx")
    sys.exit(1)


def parse_markdown_to_docx(md_file, output_file):
    """
    Convert a markdown file to a Word document.
    
    Args:
        md_file: Path to the input markdown file
        output_file: Path to the output Word document
    """
    print(f"Converting {md_file} to {output_file}...")
    
    # Create a new Document
    doc = Document()
    
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    in_table = False
    table_data = []
    table_headers = []
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            if not in_table:
                doc.add_paragraph()
            i += 1
            continue
        
        # Handle headers
        if line.startswith('#'):
            if in_table:
                add_table_to_doc(doc, table_headers, table_data)
                in_table = False
                table_data = []
                table_headers = []
            
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            if level == 1:
                heading = doc.add_heading(text, level=1)
            elif level == 2:
                heading = doc.add_heading(text, level=2)
            elif level == 3:
                heading = doc.add_heading(text, level=3)
            else:
                heading = doc.add_heading(text, level=4)
        
        # Handle horizontal rules
        elif line.strip().startswith('---') or line.strip().startswith('***'):
            if in_table:
                add_table_to_doc(doc, table_headers, table_data)
                in_table = False
                table_data = []
                table_headers = []
            doc.add_paragraph('_' * 50)
        
        # Handle tables
        elif '|' in line and not line.strip().startswith('EOF'):
            if not in_table:
                in_table = True
                # This is the header row
                table_headers = [cell.strip() for cell in line.split('|') if cell.strip()]
            elif line.strip().startswith('|---') or line.strip().startswith('|-'):
                # This is the separator row, skip it
                pass
            else:
                # This is a data row
                row_data = [cell.strip() for cell in line.split('|') if cell.strip()]
                if row_data and not all(cell.startswith('-') for cell in row_data):
                    table_data.append(row_data)
        
        # Handle bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            if in_table:
                add_table_to_doc(doc, table_headers, table_data)
                in_table = False
                table_data = []
                table_headers = []
            
            text = line.strip()[2:]
            text = format_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.', line.strip()):
            if in_table:
                add_table_to_doc(doc, table_headers, table_data)
                in_table = False
                table_data = []
                table_headers = []
            
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = format_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle regular paragraphs
        else:
            if in_table:
                i += 1
                continue
            
            text = format_inline_markdown(line.strip())
            p = doc.add_paragraph(text)
        
        i += 1
    
    # Add any remaining table
    if in_table and table_data:
        add_table_to_doc(doc, table_headers, table_data)
    
    # Save the document
    doc.save(output_file)
    print(f"✓ Successfully created {output_file}")


def add_table_to_doc(doc, headers, data):
    """Add a table to the document."""
    if not headers or not data:
        return
    
    # Create table with header row
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = format_inline_markdown(header, for_table=True)
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for row_data in data:
        if len(row_data) != len(headers):
            # Pad or trim to match header count
            while len(row_data) < len(headers):
                row_data.append('')
            row_data = row_data[:len(headers)]
        
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = format_inline_markdown(cell_data, for_table=True)


def format_inline_markdown(text, for_table=False):
    """
    Format inline markdown elements (bold, italic, etc.).
    For tables, we return plain text since docx table cells need different handling.
    """
    if for_table:
        # For tables, just remove markdown and return plain text
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
        text = re.sub(r'`(.+?)`', r'\1', text)  # Code
        return text
    
    # For regular text, keep the markdown indicators
    # (docx library handles some of this differently)
    return text


def convert_file(input_path, output_dir=None):
    """Convert a single markdown file to Word."""
    input_path = Path(input_path)
    
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        return False
    
    # Determine output path
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_file = output_dir / (input_path.stem + '.docx')
    else:
        output_file = input_path.with_suffix('.docx')
    
    try:
        parse_markdown_to_docx(str(input_path), str(output_file))
        return True
    except Exception as e:
        print(f"Error converting {input_path}: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to Word documents for grant applications.'
    )
    parser.add_argument(
        'files',
        nargs='*',
        help='Markdown files to convert (or use --all for all files in grant_application)'
    )
    parser.add_argument(
        '--all',
        action='store_true',
        help='Convert all markdown files in grant_application directory'
    )
    parser.add_argument(
        '--output-dir',
        '-o',
        help='Output directory for Word documents (default: same as input files)'
    )
    
    args = parser.parse_args()
    
    # Get list of files to convert
    files_to_convert = []
    
    if args.all:
        grant_app_dir = Path('grant_application')
        if grant_app_dir.exists():
            files_to_convert = list(grant_app_dir.glob('*.md'))
            # Exclude README.md as it's documentation
            files_to_convert = [f for f in files_to_convert if f.name != 'README.md']
        else:
            print("Error: grant_application directory not found")
            return 1
    elif args.files:
        files_to_convert = [Path(f) for f in args.files]
    else:
        parser.print_help()
        return 1
    
    if not files_to_convert:
        print("No files to convert")
        return 1
    
    print(f"\nConverting {len(files_to_convert)} file(s) to Word format...\n")
    
    success_count = 0
    for file_path in files_to_convert:
        if convert_file(file_path, args.output_dir):
            success_count += 1
    
    print(f"\n{'='*60}")
    print(f"Conversion complete: {success_count}/{len(files_to_convert)} files converted successfully")
    print(f"{'='*60}\n")
    
    return 0 if success_count == len(files_to_convert) else 1


if __name__ == '__main__':
    sys.exit(main())
