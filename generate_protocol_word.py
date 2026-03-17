#!/usr/bin/env python3
"""
generate_protocol_word.py
=========================
Generates the filled CER-VD ethics protocol Word document (BASEC-ready)
from the swissethics "Further use with consent" template.

Output:
    grant_application/protocole_CER-VD_CVRF-CardioProtect.docx

Usage:
    python3 generate_protocol_word.py

Requirements:
    pip install python-docx
"""

import os
import sys
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "251126_protokolltemplate_weiterverwendung_einw_en.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "grant_application", "protocole_CER-VD_CVRF-CardioProtect.docx")


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------
def add_heading(doc, text, level=1):
    """Add a heading paragraph."""
    para = doc.add_heading(text, level=level)
    return para


def add_para(doc, text, bold=False, italic=False, style="Normal"):
    """Add a normal paragraph with optional formatting."""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def add_bullet(doc, text):
    """Add a bullet list item."""
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(text)
    return para


def add_table_2col(doc, rows_data, col_widths=(2.2, 4.3)):
    """Add a simple 2-column table. rows_data is list of (col1, col2) tuples."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for i, (c1, c2) in enumerate(rows_data):
        table.rows[i].cells[0].text = str(c1)
        table.rows[i].cells[1].text = str(c2)
        table.rows[i].cells[0].width = Inches(col_widths[0])
        table.rows[i].cells[1].width = Inches(col_widths[1])
    return table


def add_signature_block(doc, role, name):
    """Add a signature block with place/date and signature lines."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"{role}:").bold = True
    add_para(doc, name)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Place, date:"
    table.rows[0].cells[1].text = "Signature:"
    doc.add_paragraph()


def page_break(doc):
    """Insert a page break."""
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Main document construction
# ---------------------------------------------------------------------------
def build_protocol():
    doc = Document()

    # -----------------------------------------------------------------------
    # Cover / Title page
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("RESEARCH PLAN / PROTOCOL")
    run.bold = True
    run.font.size = Pt(16)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_para.add_run(
        "Further Use of Health-Related Personal Data for Research\n"
        "Pursuant to Articles 32 and 33 HRA"
    )
    run2.font.size = Pt(12)

    doc.add_paragraph()
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_para.add_run("Protocol version: 1.0    |    Protocol date: 17 March 2026")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Project title
    # -----------------------------------------------------------------------
    add_heading(doc, "Title of the Research Project", level=2)
    title_block = doc.add_paragraph()
    run_t = title_block.add_run(
        "CVRF \u2013 CardioProtect: Cardiovascular Risk Factors, Guideline Adherence, "
        "and Cardiotoxicity in Cardio-Oncology at the Lausanne University Hospital (CHUV)"
    )
    run_t.bold = True
    run_t.font.size = Pt(13)

    doc.add_paragraph()
    sub_title = doc.add_paragraph()
    run_st = sub_title.add_run(
        "Including the CVRF-HYPERCHOL Sub-study: Hypercholesterolaemia Prevalence, "
        "LDL-C Target Achievement, and Impact on Cardiotoxicity in Cancer Patients"
    )
    run_st.italic = True

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Project leader
    # -----------------------------------------------------------------------
    add_heading(doc, "Name and Address of the Project Leader", level=2)
    leader_info = [
        ("Name / Title:", "Dr Sarah Hugelshofer"),
        ("Position:", "Cardiologist, Cardio-Oncology Outpatient Clinic"),
        ("Department:", "Department of Cardiology"),
        ("Institution:", "Lausanne University Hospital (CHUV)"),
        ("Address:", "Rue du Bugnon 46, 1011 Lausanne, Switzerland"),
        ("Telephone:", "+41 21 314 11 11"),
        ("E-mail:", "sarah.hugelshofer@chuv.ch"),
    ]
    add_table_2col(doc, leader_info, col_widths=(1.8, 4.7))

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Sponsor
    # -----------------------------------------------------------------------
    add_heading(doc, "Name and Address of the Sponsor", level=2)
    sponsor_info = [
        ("Name / Title:", "Prof. Dr Olivier M\u00fcller"),
        ("Position:", "Head, Department of Cardiology; Full Professor, Faculty of Biology and Medicine"),
        ("Institution:", "Lausanne University Hospital (CHUV) / University of Lausanne"),
        ("Address:", "Rue du Bugnon 46, 1011 Lausanne, Switzerland"),
        ("Telephone:", "+41 21 314 11 11"),
        ("E-mail:", "olivier.muller@chuv.ch"),
    ]
    add_table_2col(doc, sponsor_info, col_widths=(1.8, 4.7))

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Confirmation / Signature block
    # -----------------------------------------------------------------------
    add_heading(doc, "Confirmation of the Project Leader and Sponsor", level=2)
    add_para(
        doc,
        "With my signature, I attest that all information in this protocol is correct "
        "and that I will comply with the information I have given and with national "
        "legislation, namely data protection law.",
    )
    doc.add_paragraph()
    add_signature_block(doc, "Project leader", "Dr Sarah Hugelshofer")
    add_signature_block(doc, "Sponsor", "Prof. Dr Olivier M\u00fcller")

    page_break(doc)

    # -----------------------------------------------------------------------
    # Abbreviations
    # -----------------------------------------------------------------------
    add_heading(doc, "Abbreviations", level=2)
    abbreviations = [
        ("ACC/AHA", "American College of Cardiology / American Heart Association"),
        ("BASEC", "Business Administration System for Ethics Committees (Switzerland)"),
        ("CER-VD", "Commission cantonale d\u2019\u00e9thique de la recherche sur l\u2019\u00eatre humain du canton de Vaud"),
        ("CHUV", "Centre Hospitalier Universitaire Vaudois (Lausanne University Hospital)"),
        ("CRF", "Case Report Form"),
        ("CV", "Cardiovascular"),
        ("CVRF", "Cardiovascular Risk Factors"),
        ("ECG", "Electrocardiogram"),
        ("ESC", "European Society of Cardiology"),
        ("FEVG / LVEF", "Left Ventricular Ejection Fraction"),
        ("GLS", "Global Longitudinal Strain"),
        ("HbA1c", "Glycated haemoglobin"),
        ("HDL-C", "High-density lipoprotein cholesterol"),
        ("HER2", "Human epidermal growth factor receptor 2"),
        ("HRA", "Human Research Act (Loi relative \u00e0 la recherche sur l\u2019\u00eatre humain, LRH)"),
        ("HRO", "Human Research Ordinance"),
        ("HTA", "Arterial hypertension"),
        ("IC / HF", "Heart failure (insuffisance cardiaque)"),
        ("IEC/ARA2", "ACE inhibitor / Angiotensin II receptor blocker"),
        ("LDL-C", "Low-density lipoprotein cholesterol"),
        ("Lp(a)", "Lipoprotein(a)"),
        ("MACE", "Major adverse cardiovascular events"),
        ("mTOR", "Mammalian target of rapamycin"),
        ("NT-proBNP", "N-terminal pro-brain natriuretic peptide"),
        ("PCSK9i", "Proprotein convertase subtilisin/kexin type 9 inhibitor"),
        ("PI", "Principal Investigator"),
        ("REDCap", "Research Electronic Data Capture"),
        ("RCT", "Randomised Controlled Trial"),
        ("SGLT2i", "Sodium-glucose co-transporter 2 inhibitor"),
        ("TNM", "Tumour-Node-Metastasis staging"),
        ("URC", "Unit\u00e9 de Recherche Clinique (Clinical Research Unit)"),
        ("VEGF", "Vascular endothelial growth factor"),
    ]
    add_table_2col(doc, abbreviations, col_widths=(1.5, 5.0))

    page_break(doc)

    # -----------------------------------------------------------------------
    # 1. Background
    # -----------------------------------------------------------------------
    add_heading(doc, "1. Background", level=1)

    add_heading(doc, "1.1 The Growing Challenge of Cardiovascular Toxicity in Oncology", level=2)
    add_para(
        doc,
        "Advances in oncology have transformed many cancers into chronic diseases: in Switzerland, more than "
        "60% of patients survive beyond five years from diagnosis. This progress has uncovered a critical clinical "
        "challenge: the cardiovascular toxicity of anticancer treatments. The risk of dying from cardiovascular "
        "disease is doubled in cancer survivors compared with the general population, and the probability of "
        "developing heart failure is increased by 50% (Nature, 2020).",
    )
    add_para(
        doc,
        "Anthracyclines, anti-HER2 agents, tyrosine kinase inhibitors, and immunotherapies induce cardiotoxicity "
        "ranging from subclinical dysfunction to severe heart failure. Our cardio-oncology outpatient clinic at "
        "CHUV (5\u20136 patients per day, 4\u20135 days per week since January 2025, approximately 1,000 patients "
        "per year) provides daily first-hand observation of these complications.",
    )
    add_para(
        doc,
        "According to the European Society of Cardiology (ESC) guidelines on cardio-oncology (2022), the creation "
        "of cardio-oncology registries is strongly encouraged to monitor care quality and contribute to clinical "
        "research in this evidence-poor field.",
    )

    add_heading(doc, "1.2 Cardiovascular Risk Factors in the Oncological Population", level=2)
    add_para(
        doc,
        "Cardiovascular risk factors (CVRFs) play a key role and increase the risk of both cardiovascular disease "
        "and certain cancers. Cancer patients therefore frequently present with CVRFs, with the following estimated "
        "prevalences (CARDIOTOX Registry, N > 10,000; EC3 European Study, N > 8,000; Herrmann et al., Circ Res 2022; "
        "Sturgeon et al., EHJ 2019):",
    )
    cvrf_prev = [
        ("Diabetes mellitus", "20\u201330%"),
        ("Arterial hypertension", "35\u201350%"),
        ("Dyslipidaemia / Hypercholesterolaemia", "30\u201340%"),
        ("Obesity (BMI \u226530 kg/m\u00b2)", "25\u201335%"),
        ("Established cardiovascular disease", "15\u201325%"),
    ]
    add_table_2col(doc, cvrf_prev, col_widths=(3.0, 3.5))
    doc.add_paragraph()
    add_para(
        doc,
        "These CVRFs substantially increase the risk of cardiotoxicity and strongly influence patients' overall "
        "prognosis. Patients with diabetes have a 2\u20133-fold higher risk of cardiac dysfunction under "
        "anthracyclines (Lipshultz et al., JCO 2012; Zamorano et al., EHJ 2016). Uncontrolled hypertension "
        "doubles the risk (Cardinale et al., JACC 2015; HR 2.1, p=0.003). Each additional CVRF increases "
        "the risk of cardiotoxicity by 30\u201340% (Herrmann et al., Circulation 2016).",
    )
    add_para(
        doc,
        "Two critical questions emerge from daily clinical practice: (1) Are CVRFs systematically screened "
        "in cancer patients before and during cardiotoxic treatment? (2) During cardiological follow-up, "
        "are CVRFs treated to guideline-recommended targets?",
        bold=False,
        italic=True,
    )

    add_heading(doc, "1.3 The Particular Importance of Hypercholesterolaemia", level=2)
    add_para(
        doc,
        "Among all CVRFs present in the oncological population, hypercholesterolaemia is particularly "
        "important for several reasons:",
    )
    bullets_hyperchol = [
        "High prevalence: 30\u201340% of cancer patients present with significant dyslipidaemia.",
        "Modifiable risk factor: treatable with statins, ezetimibe, and PCSK9 inhibitors.",
        "Interaction with cardiotoxicity: hypercholesterolaemia aggravates anthracycline-induced myocardial "
        "injury (HR 1.8 for cardiotoxicity; Cardinale et al., JACC 2015). Oxidised LDL particles and "
        "anthracyclines share free-radical generation mechanisms, creating a synergistic deleterious effect.",
        "Cardioprotective potential of statins: meta-analysis by Tini et al. (JACC: CardioOncology, 2022, "
        "N = 3,818): 44% reduction in cardiotoxicity risk (RR 0.56; 95%CI 0.39\u20130.81), confirmed by "
        "multiple RCTs (Acar 2011; Chotenimitkhun 2015; Nabati 2019).",
        "Undertreatment: fewer than 50% of hypercholesterolaemic cancer patients achieve ESC 2022 LDL-C "
        "targets (Alexandre et al., EHJ 2020).",
        "New lipid-lowering therapies: PCSK9 inhibitors (inclisiran, twice-yearly subcutaneous injection) "
        "are particularly well-adapted to oncology patients; no prospective data yet exist.",
    ]
    for b in bullets_hyperchol:
        add_bullet(doc, b)

    add_heading(doc, "1.4 Cardioprotective Potential of SGLT2 Inhibitors", level=2)
    add_para(
        doc,
        "Sodium-glucose co-transporter 2 inhibitors (SGLT2i), prescribed for diabetes management, "
        "possess pleiotropic myocardial properties: metabolic reprogramming, mitochondrial protection, "
        "anti-inflammatory and anti-fibrotic effects. Key evidence:",
    )
    bullets_sglt2 = [
        "Preclinical: Quagliariello et al. (2021) \u2014 empagliflozin reduces myocardial fibrosis by 60% and "
        "preserves LVEF in doxorubicin-treated mice. Luo et al. (2022) \u2014 dapagliflozin preserves "
        "mitochondrial function and reduces apoptosis by 55%.",
        "Clinical: Gongora et al. (2022) \u2014 63% reduction in cardiotoxicity in 156 patients on "
        "anthracyclines treated with SGLT2i. Abdel-Qadir et al. (JAMA Oncol, 2023, N = 4,742) \u2014 "
        "SGLT2i reduced HF hospitalisations by 35% vs other antidiabetics (HR 0.65; 95%CI 0.48\u20130.88).",
        "Ongoing trials: EMPA-CARD (N = 300, 2025) and DAPA-PROTECT (N = 500, 2026).",
        "Gap: No prospective European study in real-world conditions.",
    ]
    for b in bullets_sglt2:
        add_bullet(doc, b)

    add_heading(doc, "1.5 Cancers Sharing Risk Factors with Cardiovascular Disease", level=2)
    add_para(
        doc,
        "Cancers that share the same risk factors as metabolic syndrome represent an important sub-population "
        "within our clinic: colorectal cancer, post-menopausal breast cancer, endometrial cancer, lung cancer, "
        "renal cell carcinoma, hepatocellular carcinoma, pancreatic cancer, oesophageal adenocarcinoma, and "
        "prostate cancer. In these patients, active treatment of CVRFs \u2014 particularly hypercholesterolaemia "
        "\u2014 may reduce both cardiovascular risk and potentially influence oncological prognosis through "
        "pleiotropic effects.",
    )

    add_heading(doc, "1.6 Evidence Gaps and Justification for this Registry", level=2)
    gaps = [
        "No Swiss prospective real-world data on CVRF prevalence and management quality in cardio-oncology",
        "No data on the association between LDL-C target achievement (per ESC 2022) and cardiotoxicity "
        "in a contemporary European cohort",
        "No prospective European data on new lipid-lowering therapies (PCSK9i, inclisiran) in oncology",
        "Limited real-world data on SGLT2i and statins in European cardio-oncology populations",
        "No Swiss data on adherence to ESC 2022 cardio-oncology guidelines",
    ]
    for g in gaps:
        add_bullet(doc, g)
    doc.add_paragraph()
    add_para(
        doc,
        "This registry directly addresses these gaps. It will provide the first prospective Swiss evidence "
        "base, directly applicable to clinical practice and guideline development.",
        bold=True,
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 2. Objectives
    # -----------------------------------------------------------------------
    add_heading(doc, "2. Objectives", level=1)

    add_heading(doc, "2.1 Objectives", level=2)
    add_para(doc, "Primary Objective:", bold=True)
    add_para(
        doc,
        "To describe the characteristics of the cardio-oncology population at CHUV, including the prevalence "
        "and management quality of cardiovascular risk factors (with particular focus on hypercholesterolaemia), "
        "cardiovascular treatments prescribed, and adherence to ESC 2022 cardio-oncology guideline recommendations.",
    )
    doc.add_paragraph()
    add_para(doc, "Secondary Objectives:", bold=True)
    sec_obj = [
        "To determine the 12-month incidence of cardiotoxicity (defined by ESC 2022 criteria)",
        "To assess the quality of CVRF control (hypertension, diabetes, dyslipidaemia, obesity) during "
        "cardiological follow-up",
        "To quantify the association between LDL-C target achievement (< 1.4 mmol/L or < 1.8 mmol/L per "
        "ESC 2022 risk level) and cardiotoxicity at 12 months",
        "To evaluate therapeutic modifications induced by specialist cardiological follow-up (initiation, "
        "intensification, or substitution of cardiovascular treatments)",
        "To analyse the impact of lipid-lowering treatment intensity (low/moderate/high-intensity statin "
        "\u00b1 ezetimibe \u00b1 PCSK9 inhibitor) on cardiotoxicity",
        "To explore variations in the lipid profile (LDL-C, HDL-C, triglycerides, Lp(a)) during oncological "
        "treatment and cardiological follow-up",
        "To assess major adverse cardiovascular events (MACE), arrhythmias, treatment modifications, and "
        "all-cause mortality",
        "To evaluate adherence to prescribed cardiovascular medications",
        "To identify factors associated with failure to achieve lipid targets",
    ]
    for i, obj in enumerate(sec_obj, 1):
        add_bullet(doc, f"{i}. {obj}")

    doc.add_paragraph()
    add_para(doc, "Hypotheses:", bold=True)
    hyp = [
        "Patients with diabetes and dyslipidaemia demonstrate improved CVRF control during cardiological "
        "follow-up in the cardio-oncology clinic compared with baseline",
        "LDL-C target achievement per ESC 2022 is associated with a reduced risk of cardiotoxicity at 12 months",
        "At inclusion, fewer than 50% of hypercholesterolaemic cancer patients achieve their lipid targets",
        "Cardiological follow-up leads to a statistically significant improvement in lipid management",
    ]
    for h in hyp:
        add_bullet(doc, h)

    add_heading(doc, "2.2 Endpoints", level=2)
    add_para(doc, "Primary endpoint:", bold=True)
    add_para(
        doc,
        "Cardiotoxicity at 12 months, defined according to ESC 2022 criteria:",
    )
    add_bullet(doc, "Clinical cardiotoxicity: new symptomatic heart failure OR LVEF < 50% with a drop of "
               "\u226510 percentage points from baseline")
    add_bullet(doc, "Subclinical cardiotoxicity: reduction of Global Longitudinal Strain (GLS) > 15% from baseline")
    doc.add_paragraph()
    add_para(doc, "Secondary endpoints:", bold=True)
    endpoints = [
        ("LDL-C target achievement", "Proportion achieving LDL-C < 1.4 or < 1.8 mmol/L at 6 months per ESC 2022"),
        ("Absolute LDL-C change", "\u0394LDL-C (mmol/L) from baseline to 6 and 12 months"),
        ("MACE", "Composite: cardiovascular death, non-fatal MI, non-fatal stroke, HF hospitalisation"),
        ("Arrhythmia", "New atrial fibrillation, QTc > 500 ms, ventricular arrhythmia"),
        ("Treatment modification", "Dose reduction, interruption, or discontinuation due to cardiovascular toxicity"),
        ("All-cause mortality", "Death from any cause during follow-up"),
        ("CV hospitalisation", "Any unplanned hospitalisation for cardiovascular cause"),
        ("Biomarker elevation", "hs-TnI/T or NT-proBNP above upper limits of normal"),
        ("Treatment initiation/intensification", "New or intensified lipid-lowering therapy after cardiological evaluation"),
        ("CVRF control at 12 months", "HbA1c < 7% (diabetics); SBP < 130 mmHg (hypertensives)"),
    ]
    add_table_2col(doc, endpoints, col_widths=(2.5, 4.0))

    add_heading(doc, "2.3 Independent Variables", level=2)
    add_para(
        doc,
        "Primary variable of interest \u2014 CVRF-HYPERCHOL sub-study: LDL-C target achievement at 6 months "
        "(binary: yes/no per ESC 2022 thresholds) as the exposure variable for the cardiotoxicity outcome.",
    )
    add_para(
        doc,
        "Primary variable of interest \u2014 SGLT2i analysis: SGLT2 inhibitor use at baseline and during "
        "follow-up (binary: yes/no) in diabetic patients.",
    )
    add_para(doc, "Pre-specified confounders (adjustment variables):", bold=True)
    confounders = [
        "Age and sex",
        "Cancer type and stage (TNM)",
        "Anthracycline use (yes/no, cumulative dose in mg/m\u00b2)",
        "Anti-HER2 therapy (yes/no)",
        "Thoracic radiotherapy (yes/no, dose)",
        "Tyrosine kinase inhibitors (yes/no)",
        "Immunotherapy/checkpoint inhibitors (yes/no)",
        "Baseline LVEF (%)",
        "Number of CVRFs at baseline",
        "Renal function at baseline (eGFR, CKD-EPI formula)",
    ]
    for c in confounders:
        add_bullet(doc, c)

    page_break(doc)

    # -----------------------------------------------------------------------
    # 3. Design
    # -----------------------------------------------------------------------
    add_heading(doc, "3. Design", level=1)
    add_para(doc, "Study type:", bold=True)
    add_para(doc, "Prospective observational cohort study with a retrospective component.")
    doc.add_paragraph()
    add_para(doc, "Setting:", bold=True)
    add_para(
        doc,
        "Cardio-oncology Outpatient Clinic, Department of Cardiology, Lausanne University Hospital (CHUV), "
        "Lausanne, Switzerland. This is the reference centre for cardio-oncology in French-speaking Switzerland, "
        "seeing approximately 1,000 patients per year since January 2025.",
    )
    doc.add_paragraph()
    add_para(doc, "Study structure:", bold=True)
    add_bullet(doc, "Retrospective component: data collected January\u2013December 2025 (~500 patients already "
               "registered in REDCap) analysed after ethics committee approval.")
    add_bullet(doc, "Prospective component: new patients attending the cardio-oncology clinic from 2026 onwards "
               "enrolled prospectively (~400 patients in 2026).")
    add_bullet(doc, "Follow-up: all patients followed for up to 12 months with structured evaluations at "
               "3, 6, and 12 months.")
    doc.add_paragraph()
    add_para(doc, "Collection period:", bold=True)
    add_para(doc, "January 2025 \u2013 December 2026 (data collection); follow-up analysis through December 2027.")
    doc.add_paragraph()
    add_para(doc, "Evaluation methods and data collected:", bold=True)
    methods = [
        "Clinical cardiology examination (medical history, physical examination, blood pressure)",
        "Transthoracic echocardiography: LVEF by modified Simpson\u2019s biplane method; GLS by 2D "
        "speckle-tracking echocardiography (ASE/EACVI protocol)",
        "12-lead electrocardiography (QTc interval, rhythm analysis)",
        "Blood sampling: high-sensitivity troponin I/T, NT-proBNP, HbA1c, complete lipid profile "
        "(LDL-C, HDL-C, total cholesterol, triglycerides, non-HDL-C, Lp(a), ApoB/ApoA1), "
        "creatinine/eGFR, ALAT/ASAT, CK, TSH (where clinically indicated)",
        "Standardised electronic Case Report Forms in the CHUV REDCap registry",
        "Structured collection of oncological treatment data from medical records: type of treatment, "
        "dose (including cumulative anthracycline dose), schedule, start and end dates",
    ]
    for m in methods:
        add_bullet(doc, m)
    doc.add_paragraph()
    add_para(
        doc,
        "Note: The SPHN risk assessment tool for health-related data de-identification has been considered. "
        "Given the use of coded data, the CHUV institutional REDCap infrastructure, and the restricted access "
        "to the key, the re-identification risk is assessed as low.",
        italic=True,
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 4. Origin of the Data
    # -----------------------------------------------------------------------
    add_heading(doc, "4. Origin of the Data", level=1)
    add_para(
        doc,
        "The data source is the institutional cardio-oncology registry at CHUV, implemented in REDCap "
        "(Research Electronic Data Capture), operational since January 2025. All data originate from routine "
        "clinical consultations at the CHUV Cardio-Oncology Outpatient Clinic and were/are collected as part of "
        "standard clinical care, recorded in both the institutional electronic medical record (Soarian, CHUV) "
        "and the REDCap research registry.",
    )
    doc.add_paragraph()
    add_para(doc, "Nature of data used:", bold=True)
    data_types = [
        "Demographic data: age, sex, weight, height, smoking status",
        "Cancer data: type, histology, stage (TNM), oncological treatment received (type, dose, schedule)",
        "Cardiovascular data: medical history (prior CVD, CVRF diagnoses), echocardiographic measurements "
        "(LVEF, GLS, diastolic parameters), ECG findings, blood pressure measurements, biomarker values "
        "(troponin, NT-proBNP)",
        "Biological data: lipid profile (LDL-C, HDL-C, total cholesterol, triglycerides, Lp(a)), "
        "glycaemic parameters (HbA1c, fasting glucose), renal function (creatinine, eGFR), "
        "liver enzymes (ALAT, ASAT), CK, TSH",
        "Medication data: cardiovascular treatments (statins, PCSK9 inhibitors, SGLT2i, antihypertensives, "
        "anticoagulants), oncological treatments, comedications",
        "Outcome data: cardiovascular events, hospitalisations, treatment modifications, death",
    ]
    for d in data_types:
        add_bullet(doc, d)
    doc.add_paragraph()
    add_para(
        doc,
        "No biological material (tissue samples, blood samples, biobank specimens) is used in this study. "
        "All analyses are based exclusively on health-related personal data collected during routine clinical care.",
        bold=True,
    )
    doc.add_paragraph()
    add_para(doc, "Population studied:", bold=True)
    add_para(
        doc,
        "Adult patients (\u226518 years) with an oncological diagnosis referred for cardiotoxicity evaluation "
        "and/or monitoring at the CHUV Cardio-Oncology Outpatient Clinic. This population is not considered a "
        "vulnerable population as defined by Swiss legislation (HRA Art. 2), as all subjects are competent "
        "adults; however, the active cancer diagnosis is taken into account in the ethical risk-benefit "
        "assessment (Section 13).",
    )
    add_para(
        doc,
        "This population is directly relevant to the scientific questions: it represents the patients for "
        "whom CVRF management in the context of cardiotoxic cancer treatment is most critical. No other "
        "population would be appropriate to answer these research hypotheses.",
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 5. Inclusion Criteria
    # -----------------------------------------------------------------------
    add_heading(doc, "5. Inclusion Criteria", level=1)
    inclusions = [
        "Patient attending the CHUV Cardio-Oncology Outpatient Clinic between 1 January 2025 and "
        "31 December 2026",
        "Age \u226518 years at the time of first consultation",
        "Active oncological diagnosis (solid tumour or haematological malignancy) requiring cardiovascular "
        "evaluation and/or monitoring for cardiotoxicity",
        "At least one clinical consultation documented in the CHUV cardio-oncology REDCap registry",
        "CVRF-HYPERCHOL sub-study (additional): lipid profile available at the time of first consultation "
        "(minimum: LDL-C and total cholesterol)",
    ]
    for i, inc in enumerate(inclusions, 1):
        add_bullet(doc, f"{i}. {inc}")

    # -----------------------------------------------------------------------
    # 6. Exclusion Criteria
    # -----------------------------------------------------------------------
    add_heading(doc, "6. Exclusion Criteria", level=1)
    exclusions = [
        "Refusal to participate or documented objection under CHUV General Consent (retrospective arm) "
        "or refusal of written informed consent (prospective arm)",
        "Absence of any consultable medical record in the CHUV electronic health record system (Soarian) "
        "or in the REDCap registry",
        "Age < 18 years",
        "CVRF-HYPERCHOL sub-study (additional): complete absence of lipid data in the medical record "
        "(no LDL-C or total cholesterol measurement available at inclusion or within the 6 months preceding "
        "inclusion)",
    ]
    for i, exc in enumerate(exclusions, 1):
        add_bullet(doc, f"{i}. {exc}")

    page_break(doc)

    # -----------------------------------------------------------------------
    # 7. Informed Consent
    # -----------------------------------------------------------------------
    add_heading(doc, "7. Informed Consent and Information of Participants", level=1)

    add_para(doc, "Retrospective arm (patients consulted January\u2013December 2025):", bold=True)
    add_para(
        doc,
        "All patients seen at CHUV have the opportunity to sign the CHUV General Consent "
        "(Consentement g\u00e9n\u00e9ral), which authorises the use of routinely collected health-related "
        "data for research purposes. Data from patients who have signed the general consent will be used for "
        "the retrospective analysis. A copy of the CHUV General Consent form is attached as an annex to "
        "this protocol.",
    )
    add_para(
        doc,
        "For patients who have not signed the general consent, the right-of-objection procedure (per "
        "HRA Art. 32) will be applied: patients will be individually informed of the research project "
        "and given the opportunity to object to the use of their data.",
    )
    doc.add_paragraph()
    add_para(doc, "Prospective arm (patients from 2026 onwards):", bold=True)
    add_para(
        doc,
        "All newly enrolled patients will be provided with a written Participant Information Sheet and "
        "Informed Consent Form (PICF), approved by the CER-VD, before any study-specific data collection. "
        "Participation is entirely voluntary. Patients who decline participation continue to receive full "
        "standard clinical care without any disadvantage.",
    )
    add_para(
        doc,
        "The treating cardiologist or an authorised study nurse will explain the study to the patient, "
        "provide the PICF, allow adequate time for questions, and obtain written consent before inclusion. "
        "The original signed consent form will be kept in the CHUV clinical research archive, separate from "
        "study data.",
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 8. Scientific Methods and Sample Size
    # -----------------------------------------------------------------------
    add_heading(doc, "8. Scientific Methods and Sample Size", level=1)

    add_heading(doc, "8.1 Sample Size", level=2)
    add_para(doc, "Total cohort:", bold=True)
    add_para(
        doc,
        "Approximately 850\u2013900 patients over the study period (January 2025 \u2013 December 2026):",
    )
    add_bullet(doc, "Retrospective component: ~500 patients (2025 data already collected in REDCap)")
    add_bullet(doc, "Prospective component: ~350\u2013400 patients (new enrolments in 2026)")
    doc.add_paragraph()
    add_para(doc, "CVRF-HYPERCHOL sub-study:", bold=True)
    add_para(
        doc,
        "Approximately 250\u2013350 patients with documented hypercholesterolaemia (estimated prevalence "
        "30\u201340% of the total cohort).",
    )
    doc.add_paragraph()
    add_para(doc, "Justification:", bold=True)
    add_bullet(doc, "Estimated 12-month incidence of cardiotoxicity: 12\u201315% (~100\u2013130 events "
               "in the total cohort)")
    add_bullet(doc, "With ~850 patients and ~100\u2013130 events, >80% power to detect a relative risk "
               "reduction of 40\u201350% after multivariate adjustment (two-sided alpha = 0.05)")
    add_bullet(doc, "HYPERCHOL sub-study: ~30\u201350 events in ~300 hypercholesterolaemic patients, "
               "sufficient for the primary analysis")

    add_heading(doc, "8.2 Statistical Methods", level=2)
    add_para(doc, "Descriptive statistics:", bold=True)
    add_bullet(doc, "Continuous variables: mean \u00b1 SD or median (IQR) according to distribution")
    add_bullet(doc, "Categorical variables: frequencies and percentages with 95% confidence intervals")
    add_bullet(doc, "Prevalence of CVRFs at baseline (with 95%CI)")
    add_bullet(doc, "Distribution of lipid profiles and proportion achieving ESC 2022 LDL-C targets")

    doc.add_paragraph()
    add_para(doc, "Analyses for primary endpoint (cardiotoxicity at 12 months):", bold=True)
    add_bullet(doc, "Logistic regression (unadjusted and multivariate): dependent variable = cardiotoxicity "
               "(yes/no); primary exposure = LDL-C target achievement (HYPERCHOL) or SGLT2i use (diabetics); "
               "adjustment for pre-specified confounders (Section 2.3)")
    add_bullet(doc, "Propensity score matching (PSM) to reduce selection bias in treated vs untreated "
               "comparisons (statins, SGLT2i)")

    doc.add_paragraph()
    add_para(doc, "Analyses for secondary endpoints:", bold=True)
    add_bullet(doc, "Time-to-event analyses: Kaplan-Meier curves, log-rank test, Cox proportional hazards "
               "models for time to cardiotoxicity and MACE")
    add_bullet(doc, "Longitudinal analyses: linear mixed models for repeated echocardiographic and "
               "biological measurements (LVEF, GLS, LDL-C, HbA1c) over 3, 6, and 12 months")
    add_bullet(doc, "Mediation analysis: decomposing cardioprotective effect of statins into LDL-C-mediated "
               "vs pleiotropic components")
    add_bullet(doc, "Subgroup analyses (pre-specified): by cancer type, anthracycline use, LDL-C "
               "quartiles at baseline, presence of additional CVRFs, statin intensity")

    doc.add_paragraph()
    add_para(doc, "Missing data: Multiple imputation by chained equations (MICE) for missing secondary "
             "variables; complete case analysis for primary endpoint.", italic=True)
    add_para(doc, "Significance level: p < 0.05 (two-sided) for all pre-specified analyses.", italic=True)
    add_para(doc, "Software: R (\u226504.3) and Stata (\u226518). Analyses performed by a qualified "
             "biostatistician from the CHUV Clinical Research Unit.", italic=True)

    add_heading(doc, "8.3 Potential Limitations and Risk of Bias", level=2)
    limitations = [
        "Selection bias: the cardio-oncology clinic receives referred patients who may over-represent "
        "higher cardiovascular risk individuals. Comparison with general oncology populations will be performed.",
        "Observational design: causal inferences cannot be definitively established. Propensity score "
        "methods and sensitivity analyses will be used.",
        "Loss to follow-up: competing risk analyses will be performed where appropriate.",
        "Lipid data variability: oncological treatments (corticosteroids, mTOR inhibitors, ADT) may "
        "themselves alter lipid profiles; these will be documented and controlled for in sensitivity analyses.",
        "Echocardiographic reproducibility: all measurements will follow standardised ASE/EACVI protocols; "
        "inter-observer variability will be assessed.",
    ]
    for lim in limitations:
        add_bullet(doc, lim)

    page_break(doc)

    # -----------------------------------------------------------------------
    # 9. Reporting Obligations
    # -----------------------------------------------------------------------
    add_heading(doc, "9. Reporting Obligations", level=1)
    add_para(
        doc,
        "The ethics committee (CER-VD) is notified of any change of project leader in advance.",
    )
    add_para(
        doc,
        "The project leader notifies the ethics committee of the completion or early termination of the "
        "research project within 90 days of the end or termination date.",
    )
    add_para(
        doc,
        "Any substantial amendments to the protocol (changes to objectives, design, population, "
        "statistical methods, or data protection measures) will be submitted to the CER-VD for approval "
        "before implementation.",
    )
    add_para(
        doc,
        "A progress report will be submitted to the CER-VD at 12 months. Annual progress reports will "
        "be submitted to all funding bodies.",
    )

    # -----------------------------------------------------------------------
    # 10. Data Protection
    # -----------------------------------------------------------------------
    add_heading(doc, "10. Data Protection", level=1)
    add_para(doc, "Coding of data:", bold=True)
    add_para(
        doc,
        "All health-related personal data used in this research project are coded (pseudo-anonymised). "
        "Each patient is assigned a unique anonymous participant identification number (e.g., CVRF-001, "
        "CVRF-002, \u2026). This code replaces all direct identifiers (name, date of birth, hospital record "
        "number, address) in the research database.",
    )
    doc.add_paragraph()
    add_para(doc, "Key management:", bold=True)
    add_para(
        doc,
        "Project data will be handled with utmost discretion and are only accessible to authorised personnel "
        "who require the data to fulfil their duties within the scope of the research project. On Case Report "
        "Forms, participants are identified only by their participant ID.",
    )
    add_para(
        doc,
        "The coding key (linking participant ID to real patient identifiers) is held exclusively by "
        "Dr Sarah Hugelshofer (project leader) and is stored in a password-protected, encrypted file on the "
        "CHUV secured institutional server, separate from the REDCap research database. Access is restricted "
        "to Dr Hugelshofer and, under her supervision, authorised CHUV research staff. The key is never "
        "transferred outside CHUV.",
    )
    doc.add_paragraph()
    add_para(doc, "Data platform:", bold=True)
    add_para(
        doc,
        "The REDCap institutional installation at CHUV is used as the research database. REDCap is hosted "
        "on CHUV-managed servers located in Switzerland. The platform is ISO 27001-certified and compliant "
        "with the Swiss Data Protection Act (nDSG/LPD) and GDPR. Role-based access controls ensure each "
        "team member accesses only the data necessary for their function. Audit trail functionality is enabled.",
    )
    doc.add_paragraph()
    add_para(doc, "Data transfer:", bold=True)
    add_para(
        doc,
        "No transfer of identifiable or coded research data outside CHUV or Switzerland is planned. "
        "Statistical analyses will be performed on de-identified aggregated datasets on CHUV secure "
        "infrastructure.",
    )
    doc.add_paragraph()
    add_para(doc, "Technical and organisational measures:", bold=True)
    measures = [
        "All digital documents are password-protected",
        "Paper CRFs (if used) are stored in locked cabinets within the CHUV Department of Cardiology "
        "research office",
        "Access to REDCap is controlled by user-level authentication (individual username and password)",
        "Audit trail functionality is enabled (all data entries and modifications logged)",
        "Regular encrypted backups on CHUV institutional servers",
    ]
    for m in measures:
        add_bullet(doc, m)

    page_break(doc)

    # -----------------------------------------------------------------------
    # 11. Storage
    # -----------------------------------------------------------------------
    add_heading(doc, "11. Information on the Storage of Data and Completion of the Research Project", level=1)
    add_para(doc, "Storage platform and location:", bold=True)
    add_para(
        doc,
        "All research data are stored exclusively in the CHUV institutional REDCap installation, on servers "
        "located within Switzerland (CHUV data centre, Lausanne). No cloud services based outside Switzerland "
        "or the European Economic Area are used.",
    )
    doc.add_paragraph()
    add_para(doc, "Data security framework:", bold=True)
    security = [
        "REDCap institutional CHUV installation with ISO 27001 certification",
        "CHUV IT security standards apply (annual security audits, encrypted connections, role-based "
        "access control)",
        "Encrypted backup on separate CHUV servers with daily backup cycles",
        "Physical access to servers controlled by CHUV IT department",
        "Paper documents: locked filing cabinets, restricted access",
    ]
    for s in security:
        add_bullet(doc, s)
    doc.add_paragraph()
    add_para(doc, "At project completion:", bold=True)
    add_bullet(doc, "Research data in REDCap will be anonymised (irreversible removal of participant ID "
               "by destruction of the coding key) or archived in anonymised form on CHUV institutional servers")
    add_bullet(doc, "The coding key will be destroyed by the project leader")
    add_bullet(doc, "Paper documents containing personal data will be securely shredded")
    add_bullet(doc, "Anonymised research data will be retained for the required statutory period (see "
               "Section 12) and may subsequently be shared as anonymised datasets in accordance with "
               "open science principles")

    # -----------------------------------------------------------------------
    # 12. Retention Period
    # -----------------------------------------------------------------------
    add_heading(doc, "12. Retention Period", level=1)
    add_para(
        doc,
        "In accordance with Swiss law (HRA Art. 47 and HRO Art. 60) and CHUV institutional policy:",
    )
    retention = [
        "Health-related research data: retained for a minimum of 10 years after the end of the research "
        "project (i.e., until at least 31 December 2037)",
        "Signed informed consent forms: retained for a minimum of 10 years after the end of the research "
        "project",
        "Coding key: retained during the active research project only, then destroyed upon data "
        "anonymisation at project completion",
    ]
    for r in retention:
        add_bullet(doc, r)
    doc.add_paragraph()
    add_para(
        doc,
        "Storage location during retention period: CHUV institutional secure servers (Lausanne, "
        "Switzerland) and CHUV Department of Cardiology research archive (paper documents).",
    )
    add_para(
        doc,
        "Responsible person: Dr Sarah Hugelshofer (or her institutional successor at CHUV) is responsible "
        "for ensuring the correct retention and eventual destruction/anonymisation of research data.",
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 13. Ethical and Regulatory Requirements
    # -----------------------------------------------------------------------
    add_heading(doc, "13. Ethical and Regulatory Requirements", level=1)

    add_heading(doc, "13.1 Risk-Benefit Assessment", level=2)
    add_para(doc, "Expected benefits:", bold=True)
    benefits = [
        "First prospective Swiss registry in cardio-oncology with comprehensive CVRF characterisation",
        "First Swiss data on LDL-C target achievement and its association with cardiotoxicity",
        "Validation (or refutation) of preclinical and small-RCT evidence in real-world conditions",
        "Direct contribution to ESC guideline development in cardio-oncology",
        "Expected 5 publications in high-impact journals (JACC: CardioOncology, European Heart Journal, "
        "Diabetes Care); expected cumulative impact factor 80\u2013150",
        "Improved CVRF management for cancer patients at CHUV through structured monitoring",
        "Foundation for future RCTs, multicentre extension, and PCSK9i studies in oncology",
        "Training of future cardiologists in evidence-based cardio-oncology",
    ]
    for b in benefits:
        add_bullet(doc, b)
    doc.add_paragraph()
    add_para(doc, "Data risks and their prevention:", bold=True)
    risks = [
        "Risk of re-identification: low, mitigated by strict pseudo-anonymisation, secure key management "
        "(Section 10), restricted data access, and CHUV-hosted infrastructure",
        "Risk of data breach: mitigated by REDCap\u2019s institutional security framework, encrypted "
        "storage, and password-protected access (Section 11)",
        "Risk of stigmatisation/discrimination: minimal; no genetic data, psychiatric diagnoses, or "
        "socially sensitive data beyond common medical conditions",
    ]
    for r in risks:
        add_bullet(doc, r)
    doc.add_paragraph()
    add_para(doc, "Harm assessment:", bold=True)
    add_para(
        doc,
        "The risk of harm to participants is considered minimal: this is a purely observational study; "
        "no invasive procedures, additional blood sampling, or additional imaging are required beyond "
        "standard clinical care. The expected scientific and clinical benefits substantially outweigh "
        "the minimal risks.",
    )
    doc.add_paragraph()
    add_para(doc, "Liability:", bold=True)
    add_para(
        doc,
        "The project leader (Dr Sarah Hugelshofer) and the institution (CHUV) bear responsibility for all "
        "aspects of data handling and data security within the scope of this protocol. No separate research "
        "insurance is required for this type of observational study.",
    )

    add_heading(doc, "13.2 Regulatory Compliance", level=2)
    add_para(
        doc,
        "This project complies with the regulatory requirements of the Swiss Human Research Act "
        "(HRA, SR 810.30) and the Human Research Ordinance (HRO, SR 810.301). The prerequisite for "
        "carrying out the research project is the approval of the Commission cantonale d\u2019\u00e9thique "
        "de la recherche sur l\u2019\u00eatre humain du canton de Vaud (CER-VD).",
    )
    add_para(
        doc,
        "The project also complies with the Swiss Data Protection Act (nDSG/LPD, effective "
        "1 September 2023), CHUV institutional research governance regulations, and ICH-GCP principles "
        "as applicable to observational studies.",
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 14. Results / Transparency / Publication
    # -----------------------------------------------------------------------
    add_heading(doc, "14. Results / Transparency / Publication", level=1)
    add_para(
        doc,
        "Results will be scientifically valid and generalisable to comparable European cardio-oncology "
        "populations, with the limitation that this is a single-centre Swiss study. The prospective design, "
        "standardised data collection, and pre-specified analysis plan minimise risk of bias. All "
        "limitations will be thoroughly discussed in publications.",
    )
    doc.add_paragraph()
    add_para(doc, "Planned publications:", bold=True)
    publications = [
        ("1", "Study protocol", "Trials or BMC Cardiovascular Disorders", "2026 (Year 1)"),
        ("2", "CVRF prevalence, management quality, guideline adherence, and cardiotoxicity at 12 months",
         "JACC: CardioOncology or European Heart Journal", "2027 (Year 2)"),
        ("3", "HYPERCHOL sub-study: LDL-C targets, statin intensity, and cardiotoxicity",
         "JACC: CardioOncology or Atherosclerosis", "2027 (Year 2)"),
        ("4", "SGLT2i sub-analysis: SGLT2 inhibitors and cardiotoxicity in diabetic cancer patients",
         "Diabetes Care", "2027\u20132028"),
        ("5", "Guideline adherence and quality of care in Swiss cardio-oncology",
         "EHJ \u2013 Quality of Care and Clinical Outcomes", "2028 (Year 3)"),
    ]
    # Use a 4-column table for publications
    table = doc.add_table(rows=len(publications) + 1, cols=4)
    table.style = "Table Grid"
    headers = ["#", "Focus / Title", "Target journal", "Timeline"]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (num, focus, journal, timeline) in enumerate(publications, 1):
        table.rows[i].cells[0].text = num
        table.rows[i].cells[1].text = focus
        table.rows[i].cells[2].text = journal
        table.rows[i].cells[3].text = timeline

    doc.add_paragraph()
    add_para(
        doc,
        "All publications will be submitted in open access format. Article processing charges are covered "
        "by research funding. Anonymised aggregate datasets supporting publications will be made available "
        "upon reasonable request after publication.",
    )
    add_para(
        doc,
        "Individual participants will not be contacted with their personal results, given the observational "
        "nature of the study. Clinically significant incidental findings identified during routine follow-up "
        "will be communicated to the treating physician per CHUV standard procedures.",
    )
    add_para(
        doc,
        "The study protocol will be registered in a publicly accessible clinical trial registry "
        "(ClinicalTrials.gov or ISRCTN) before the start of data analysis.",
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 15. Funding / Data Sharing / Declaration of Interest
    # -----------------------------------------------------------------------
    add_heading(doc, "15. Funding / Data Sharing / Declaration of Interest", level=1)
    add_para(doc, "Funding sources:", bold=True)
    funding = [
        ("Source", "Type", "Amount", "Coverage"),
        ("Fondation Emma Muschamp (CHUV)", "External grant (applied/approved)",
         "CHF 50,000", "Research associate salary (MA3 30% FTE), statistics, dissemination"),
        ("External complementary funding (obtained)", "External grant",
         "CHF 20,000", "Protected time for PI (Dr Hugelshofer, 20% FTE, 12 months)"),
        ("CHUV institutional (in-kind)", "Institutional support",
         "CHF 25,000", "Senior biostatistician (URC), REDCap/IT infrastructure, administrative support"),
        ("Novartis Pharma Schweiz AG", "Industry grant (pending)", "TBD",
         "CVRF-HYPERCHOL sub-study analyses (2026 baseline, seed funding)"),
        ("TOTAL \u2013 Phase 1", "", "CHF 95,000", ""),
    ]
    fund_table = doc.add_table(rows=len(funding), cols=4)
    fund_table.style = "Table Grid"
    for i, row_data in enumerate(funding):
        for j, cell_text in enumerate(row_data):
            fund_table.rows[i].cells[j].text = cell_text
            if i == 0:
                for run in fund_table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()
    add_para(
        doc,
        "Funding for Phase 2 (months 13\u201324) and Phase 3 (months 25\u201336) will be sought from "
        "the Swiss National Science Foundation (SNSF/FNS), the Swiss Heart Foundation, and other "
        "appropriate funding bodies.",
    )
    doc.add_paragraph()
    add_para(doc, "Publication policy:", bold=True)
    add_para(
        doc,
        "Funding sources have no right to influence study design, data collection, analysis, "
        "interpretation, or publication decisions. Scientific independence is guaranteed. All publications "
        "will acknowledge all funding sources.",
    )
    doc.add_paragraph()
    add_para(doc, "Conflicts of interest:", bold=True)
    add_bullet(doc, "Dr Sarah Hugelshofer: no conflicts of interest declared.")
    add_bullet(doc, "Prof. Olivier M\u00fcller: no conflicts of interest declared.")
    add_bullet(doc, "If Novartis Pharma Schweiz AG provides funding for the HYPERCHOL sub-study, this will "
               "be disclosed in all resulting publications. Novartis will have no influence on study design, "
               "data access, analysis, or publication decisions.")
    doc.add_paragraph()
    add_para(doc, "Data sharing / consortium agreements:", bold=True)
    add_para(
        doc,
        "This is a monocentric study. No data transfer or sharing agreements with external partners are "
        "in place at the time of protocol submission. Should multicentre extension be pursued (Phase 3), "
        "appropriate Data Transfer and Use Agreements (DTUA) will be established and submitted to "
        "the CER-VD.",
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 16. References
    # -----------------------------------------------------------------------
    add_heading(doc, "16. References", level=1)
    references = [
        "Lyon AR, et al. 2022 ESC Guidelines on cardio-oncology. Eur Heart J. 2022;43(41):4229-4361.",
        "Zamorano JL, et al. 2016 ESC Position Paper on cancer treatments and cardiovascular toxicity. "
        "Eur Heart J. 2016;37(36):2768-2801.",
        "Mach F, et al. 2019 ESC/EAS Guidelines for the management of dyslipidaemias. Eur Heart J. "
        "2020;41(1):111-188.",
        "Tini G, et al. Statin therapy and risk of cancer therapy-associated cardiotoxicity. JACC: "
        "CardioOncology. 2022;4(3):392-401.",
        "Abdel-Qadir H, et al. SGLT2 inhibitors and heart failure outcomes in cancer patients with "
        "diabetes. JAMA Oncol. 2023;9(7):950-958.",
        "Cardinale D, et al. Incidence of anthracycline cardiotoxicity and multivariable analysis of "
        "risk factors. JACC. 2015;65(25):2750-2760.",
        "Herrmann J, et al. Cardiovascular disease risk factors in patients with cancer. Circ Res. "
        "2022;130(8):1171-1182.",
        "Armenian SH, et al. Prevention of cardiovascular disease in cancer survivors. JACC: "
        "CardioOncology. 2023;5(1):1-18.",
        "Meijers WC, de Boer RA. Common risk factors for heart failure and cancer. Cardiovasc Res. "
        "2019;115(5):844-853.",
        "Koene RJ, et al. Shared risk factors in cardiovascular disease and cancer. Circulation. "
        "2016;133(11):1104-1114.",
        "Alexandre J, et al. Cardiovascular toxicity related to cancer treatment: a pragmatic perspective. "
        "Eur Heart J. 2020;41(18):1770-1779.",
        "Sturgeon KM, et al. A population-based study of cardiovascular disease mortality risk in US "
        "cancer patients. Eur Heart J. 2019;40(48):3889-3897.",
        "Gongora CA, et al. SGLT2 inhibition with empagliflozin attenuates myocardial oxidative stress "
        "and fibrosis in diabetic mice. JACC: Basic Transl Sci. 2022;7(7):661-672.",
        "Quagliariello V, et al. SGLT-2 inhibitors reduce cardiotoxicity of doxorubicin through "
        "mechanisms related to mitochondrial function and autophagy. J Clin Med. 2021;10(23):5757.",
        "Acar Z, et al. Atorvastatin attenuates cardiotoxicity induced by anthracyclines. Anadolu "
        "Kardiyol Derg. 2011;11(7):564-569.",
        "Chotenimitkhun R, et al. Chronic statin administration may attenuate early anthracycline-"
        "associated declines in left ventricular ejection function. Can J Cardiol. 2015;31(3):302-307.",
        "Nabati M, et al. Role of rosuvastatin in preventing chemotherapy-induced cardiotoxicity in "
        "women with breast cancer. J Cardiovasc Pharmacol Ther. 2019;24(3):233-241.",
        "Ligibel JA, et al. American Society of Clinical Oncology Obesity Initiative. J Clin Oncol. "
        "2022;40(23):2501-2509.",
        "Gallagher EJ, LeRoith D. Obesity and diabetes: The increased risk of cancer and cancer-related "
        "mortality. Physiol Rev. 2015;95(3):727-748.",
        "Lin Y, et al. Inhibition of PCSK9 attenuates doxorubicin-induced cardiotoxicity via "
        "suppression of oxidative stress. Cardiovasc Res. 2022;118(14):2913-2925.",
        "Luo Y, et al. Dapagliflozin protects against diabetic cardiomyopathy in type II diabetic rats. "
        "J Cardiovasc Pharmacol Ther. 2022;27:10742484221076239.",
        "Cardinale D, et al. Prevention of high-dose chemotherapy-induced cardiotoxicity by ACE "
        "inhibition. Circulation. 2006;114(23):2474-2481.",
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref)

    page_break(doc)

    # -----------------------------------------------------------------------
    # Final footer note
    # -----------------------------------------------------------------------
    add_para(doc, "Document prepared by: Dr Sarah Hugelshofer", bold=True)
    add_para(doc, "Institution: Department of Cardiology, CHUV")
    add_para(doc, "Date: 17 March 2026 | Protocol version: 1.0")
    doc.add_paragraph()
    add_para(
        doc,
        "Contact for queries: Dr Sarah Hugelshofer, Department of Cardiology, CHUV, "
        "Rue du Bugnon 46, 1011 Lausanne, Switzerland. "
        "E-mail: sarah.hugelshofer@chuv.ch. Tel.: +41 21 314 11 11",
    )
    doc.add_paragraph()
    add_para(
        doc,
        "This protocol must be submitted to the CER-VD via the BASEC web portal "
        "(https://submissions.swissethics.ch/en/). The signed PDF version (OCR format) is "
        "the official submission document. Signature pages must be printed, signed by "
        "Dr Hugelshofer and Prof. M\u00fcller, scanned in OCR format, and uploaded to BASEC.",
        italic=True,
    )

    # -----------------------------------------------------------------------
    # Save document
    # -----------------------------------------------------------------------
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"\u2705  Protocol Word document generated successfully:")
    print(f"    {OUTPUT_PATH}")


if __name__ == "__main__":
    build_protocol()
